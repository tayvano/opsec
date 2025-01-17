import os
import jinja2
from slugify import slugify
from yaml import safe_load as load
import yaml


from docx import Document
from docx.shared import Pt

from collections import OrderedDict

def ordered_load(stream, Loader=yaml.Loader, object_pairs_hook=OrderedDict):
    class OrderedLoader(Loader):
        pass
    def construct_mapping(loader, node):
        loader.flatten_mapping(node)
        return object_pairs_hook(loader.construct_pairs(node))
    OrderedLoader.add_constructor(
        yaml.resolver.BaseResolver.DEFAULT_MAPPING_TAG,
        construct_mapping)
    return yaml.load(stream, OrderedLoader)


ANSWERS = {
    "default": "Yes / No",
    "default_na": "Yes / No / Not applicable",
}

APPLIES = {
    "everyone": "Everyone",
    "mle": "Medium and large enterprises"
}


@jinja2.contextfilter
def normalize_id(jinja_ctx, context, **kw):
    return slugify(context)


def get_audit_data(chapter):
    with open("data/{}.yaml".format(chapter), "rt") as inp:
        struct = ordered_load(inp, yaml.SafeLoader)
    return struct


def get_chapters(chapter_ids):
    chapters = OrderedDict()
    for chapter_id in chapter_ids:
        with open("data/{}.yaml".format(chapter_id), "rt") as inp:
            struct = ordered_load(inp, yaml.SafeLoader)
        chapters[chapter_id] = struct
    return chapters


def get_incidences_data(chapters):
    with open("data/incidences.yaml", "rt") as inp:
        struct = ordered_load(inp, yaml.SafeLoader)

    # Build cross-reference data incidence -> question
    for incidence_id, incidence in struct.items():
        incidence["references"] = []
        for chapter_id, chapter_data in chapters.items():
            previous_question_id = None
            for question_id, question in chapter_data["questions"].items():

                try:
                    incidence_ids = question.get("incidences", [])
                except AttributeError:
                    raise RuntimeError("Bad question data: {}, previous: {}".format(question_id, previous_question_id))

                if type(incidence_ids) != list:
                    raise RuntimeError("Bad incidence list for question: {}".format(question_id))

                if incidence_id in [inc.lower() for inc in incidence_ids]:
                    incidence["references"].append((chapter_id, question_id, question["question"]))

                previous_question_id = question_id

    # Sort incidences alphabetically
    struct = OrderedDict(sorted(struct.items(), key=lambda item: item[1]["title"]))

    return struct


def friendly_number(number:str) -> int:
    """Convert 180M like number to an integer for summing."""

    multipliers = dict(M=1000000, k=1000)

    if not number:
        return 0

    if type(number) == int:
        return number

    if number[0].isdigit():
        multiplier = number[-1]
        number = number[0:-1]
        return float(number) * multipliers[multiplier]
    else:
        return 0  # "Not disclosed"


def generate_word(chapters:dict, incidences:dict):
    """Generate the word document for proof reading.

    :param chapters:
    :param incidences:
    :return:
    """

    # document.add_heading('Chapters', 0)
    # document.add_page_break()

    for chapter_id, chapter in chapters.items():

        print("Building {}.docx".format(chapter_id))
        document = Document()

        document.add_heading("Chapter: {}".format(chapter_id), 1)
        p = document.add_paragraph(chapter["meta"]["description"], style="Subtitle")

        p = document.add_paragraph()
        run = p.add_run(chapter["meta"]["background"])
        run.font.name = "Courier New"
        run.font.size = Pt(12)

        document.add_page_break()

        for question_id, question in chapter["questions"].items():
            table = document.add_table(rows=3, cols=1)

            # table.rows[0].cells[0].text = "Short title"
            table.rows[0].cells[0].text = question_id

            # table.rows[2].cells[0].text = "Title"
            table.rows[1].cells[0].text = question["question"]

            p = document.add_paragraph()
            # table.rows[4].cells[0].text = "Rationale"
            p = table.rows[2].cells[0].add_paragraph()
            try:
                run = p.add_run(question["rationale"])
            except KeyError:
                raise RuntimeError("Question missing rationale: {}".format(question_id))

            run.font.name = "Courier New"
            run.font.size = Pt(12)

            document.add_page_break()
        document.save("{}.docx".format(chapter_id))

    print("Building incidences.docx")
    document = Document()

    document.add_heading('Incidences', 0)
    document.add_page_break()

    for incidence_id, incidence in incidences.items():

        table = document.add_table(rows=3, cols=1)

        # table.rows[0].cells[0].text = "Incidence id"
        table.rows[0].cells[0].text = incidence_id

        # table.rows[2].cells[0].text = "Title"
        table.rows[1].cells[0].text = incidence["title"]

        p = document.add_paragraph()
        # table.rows[4].cells[0].text = "Description"
        p = table.rows[2].cells[0].add_paragraph()
        run = p.add_run(incidence["description"])
        run.font.name = "Courier New"
        run.font.size = Pt(12)

        document.add_page_break()

    document.save("incidences.docx".format(chapter_id))




loader = jinja2.FileSystemLoader(os.path.join(os.getcwd(), "audit_templates"))
env = jinja2.Environment(loader=loader)
env.filters["normalize_id"] = normalize_id

chapter_template = env.get_template('chapter.rst')
question_template = env.get_template('question.rst')
chapter_ids = load(open("data/index.yaml", "rt"))
chapters = get_chapters(chapter_ids)
incidences = get_incidences_data(chapters)
assets_lost = sum([friendly_number(incidence.get("assets-stolen")) for incidence in incidences.values()])
compromised_accounts = sum([friendly_number(incidence.get("compromised-users")) for incidence in incidences.values()])
incidence_count = len(incidences)
security_assesment_point_count = sum([len(chapter["questions"]) for chapter in chapters.values()])

# Generate index
index_template = env.get_template('index.rst')
doc = index_template.render(chapters=chapters, incidence_count=incidence_count, security_assesment_point_count=security_assesment_point_count)
with open("source/index.rst", "wt") as out:
    out.write(doc)



# Generate incidences
incidences_template = env.get_template('incidences.rst')


assets_lost /= 1000000;
compromised_accounts /= 1000000;

doc = incidences_template.render(incidences=incidences, assets_lost=assets_lost, compromised_accounts=compromised_accounts)
with open("source/incidences/index.rst", "wt") as out:
    out.write(doc)

# Build individual incidences pages
incidence_template = env.get_template('incidence.rst')
for incidence_id, incidence in incidences.items():
    doc = incidence_template.render(incidence=incidence, incidence_id=incidence_id)
    with open("source/incidences/{}.rst".format(incidence_id), "wt") as out:
        out.write(doc)

# Generate chapters
for chapter_id, chapter in chapters.items():
    md = chapter_template.render(answers=ANSWERS, applies=APPLIES, **chapter)

    with open("source/{}/index.rst".format(chapter_id), "wt") as out:
        out.write(md)

    for question_title, question in chapter["questions"].items():
        question_id = slugify(question_title)
        md = question_template.render(question_title=question_title, question=question, answers=ANSWERS, applies=APPLIES)
        with open("source/{}/{}.rst".format(chapter_id, question_id), "wt") as out:
            out.write(md)


generate_word(chapters, incidences)